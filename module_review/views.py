from django.shortcuts import render, get_object_or_404, redirect
from django.http import HttpResponse
from docx import Document
from docx.shared import Inches
import openpyxl
from .forms import UploadFileForm
import tempfile
import requests
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser
from drf_yasg.utils import swagger_auto_schema
from rest_framework.response import Response
from datetime import datetime, timedelta
from django.utils import timezone
from collections import Counter
from operator import itemgetter
from django.http import JsonResponse
from django.contrib import messages
from django.contrib.auth import authenticate, login
from rest_framework import status
from django.contrib.admin.views.decorators import user_passes_test
from django.contrib.auth.decorators import login_required
from django.http import HttpResponseForbidden
from account.models import User
from rest_framework.permissions import IsAuthenticated
from django.http import Http404, HttpResponseForbidden
from .models import Module, Program, ModuleReview, ProgramReview, AcademicYear
from .serializers import ModuleSerializer, ModuleReviewSerializer, ProgramSerializer, ProgramReviewSerializer, AcademicYearSerializer
from django.db.models.signals import post_save
from django.dispatch import receiver
from django.contrib.auth import get_user_model
from django.db.models import Q
from django.views.decorators.http import require_GET
# from django.contrib.auth.decorators import user_passes_test



@api_view(['GET'])
def academic_year_list(request):
    academic_years = AcademicYear.objects.all()
    serializer = AcademicYearSerializer(academic_years, many=True)
    return Response(serializer.data)


@swagger_auto_schema(method='POST', request_body=ModuleSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def module_list(request):
    if request.method == 'GET':
        user = request.user
        module_list = Module.objects.all()  # Add parentheses here
        serializer = ModuleSerializer(module_list, many=True)
        return Response(serializer.data)
    
    elif request.method == 'POST':
        data = request.data.copy() 
        data['completed_by'] = request.user.id  
        serializer = ModuleSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    
    
    
@swagger_auto_schema(method='PUT', request_body=ModuleSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
def module_detail(request, module_id):
    try:
        module_detail = Module.objects.get(module_id=module_id)
    except Module.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ModuleSerializer(module_detail)
        return Response(serializer.data)

    elif request.method == 'PUT':
        serializer = ModuleSerializer(module_detail, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        module_detail.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    
    
    
@swagger_auto_schema(method='POST', request_body=ModuleReviewSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def module_review_list(request, module_code):
    try:
        module = Module.objects.get(module_code=module_code)
    except Module.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)
    
    if request.method == 'GET':
        reviews = ModuleReview.objects.filter(module=module)
        serializer = ModuleReviewSerializer(reviews, many=True)
        return Response(serializer.data)
    
    elif request.method == 'POST':
        data = request.data.copy()
        data['module'] = module.id 
        existing_review = ModuleReview.objects.filter(module=module, academic_year=data['academic_year']).first()
        if existing_review:
            existing_review.delete()
        serializer = ModuleReviewSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    
@swagger_auto_schema(method='PUT', request_body=ModuleReviewSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def module_review_detail(request, pk):
    try:
        review = ModuleReview.objects.get(pk=pk)
    except ModuleReview.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ModuleReviewSerializer(review)
        return Response(serializer.data)
    
    elif request.method == 'PUT':
        serializer = ModuleReviewSerializer(review, data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    elif request.method == 'DELETE':
        review.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    
    
# ============================================= Program View=====================================

@swagger_auto_schema(method='POST', request_body=ProgramSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def program_list(request):
    if request.method == 'GET':
        user = request.user
        program_list = Program.objects.all()  # Add parentheses here
        serializer = ProgramSerializer(program_list, many=True)
        return Response(serializer.data)
    
    elif request.method == 'POST':
        data = request.data.copy() 
        data['completed_by'] = request.user.id  
        serializer = ProgramSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    
    
    
@swagger_auto_schema(method='PUT', request_body=ProgramSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
def program_detail(request, program_id):
    try:
        program_detail = Program.objects.get(program_id=program_id)
    except Program.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ProgramSerializer(program_detail)
        return Response(serializer.data)

    elif request.method == 'PUT':
        serializer = ProgramSerializer(program_detail, data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

    elif request.method == 'DELETE':
        program_detail.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    
    
    
@swagger_auto_schema(method='POST', request_body=ProgramReviewSerializer)
@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def program_review_list(request, program_code):
    try:
        program = Program.objects.get(program_code=program_code)
    except Program.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)
    
    if request.method == 'GET':
        reviews = ProgramReview.objects.filter(program=program)
        serializer = ProgramReviewSerializer(reviews, many=True)
        return Response(serializer.data)
    
    elif request.method == 'POST':
        data = request.data.copy()
        data['program'] = program.id    
        existing_review = ProgramReview.objects.filter(program=program, academic_year=data['academic_year']).first()
        if existing_review:
            existing_review.delete()
        serializer = ProgramReviewSerializer(data=data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


@swagger_auto_schema(method='PUT', request_body=ProgramReviewSerializer)
@api_view(['GET', 'PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def program_review_detail(request, pk):
    try:
        review = ProgramReview.objects.get(pk=pk)
    except ProgramReview.DoesNotExist:
        return Response(status=status.HTTP_404_NOT_FOUND)

    if request.method == 'GET':
        serializer = ProgramReviewSerializer(review)
        return Response(serializer.data)
    
    elif request.method == 'PUT':
        serializer = ProgramReviewSerializer(review, data=request.data, context={'request': request})
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    elif request.method == 'DELETE':
        review.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)

#=====Upload Module================================
def module_upload_xlsx(request):
    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            file = request.FILES['file']
            module_process_file(file)
            messages.success(request, "File uploaded and processed successfully")
            return redirect('admin_module_list')
    else:
        form = UploadFileForm()
    return render(request, 'admin_dashboard/upload_xlsx.html', {'form': form})

def module_process_file(file):
    wb = openpyxl.load_workbook(file)
    ws = wb.active
    for row in ws.iter_rows(min_row=2, values_only=True):
        module_code, module_name, module_leader = row
        if module_code and module_name and module_leader:
            Module.objects.update_or_create(
                module_code=module_code,
                defaults={
                    'module_name': module_name,
                    'module_leader': module_leader,
                }
            )
        
#=====Upload Program================================
def program_upload_xlsx(request):
    if request.method == "POST":
        form = UploadFileForm(request.POST, request.FILES)
        if form.is_valid():
            file = request.FILES['file']
            program_process_file(file)
            messages.success(request, "File uploaded and processed successfully")
            return redirect('admin_program_list')
    else:
        form = UploadFileForm()
    return render(request, 'admin_dashboard/upload_xlsx.html', {'form': form})


def program_process_file(file):
    wb = openpyxl.load_workbook(file)
    ws = wb.active
    for row in ws.iter_rows(min_row=2, values_only=True):
        program_name, program_leader = row
        if program_name and program_leader:
            Program.objects.update_or_create(
                program_name=program_name,
                defaults={
                    'program_leader': program_leader,
                }
            )


# ============= Export =================    
def add_row_with_spacing(table):
    row = table.add_row()
    for cell in row.cells:
        cell.paragraphs[0].add_run().add_break()
    return row

def export_module_details(request, module_id):
    module = get_object_or_404(Module, id=module_id)
    reviews = ModuleReview.objects.filter(module=module).order_by('-academic_year')
    
    document = Document()
    
    # Add logo image from URL
    logo_url = 'https://res.cloudinary.com/ddcynxxlr/image/upload/v1720968735/xypcmbpjm3szs6nwlh8e.jpg'
    response = requests.get(logo_url)
    if response.status_code == 200:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name

        header = document.sections[0].header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(tmp_file_path, width=Inches(2))  # Adjust size as needed
    
    for review in reviews:
        # Add heading for each review
        document.add_heading(f'Annual Module Quality Enhancement Report for {review.academic_year.academic_year}\n\n', level=1)
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '\nDescription\n'
        hdr_cells[1].text = '\nDetails\n'
        
        row = add_row_with_spacing(table)
        row.cells[0].text = '\n1. Module details\n'
        row.cells[1].text = f"\n{module.module_code} {module.module_name}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n2. Academic Year\n'
        row.cells[1].text = f"\n{review.academic_year.academic_year}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n3. School\n'
        row.cells[1].text = f"\n{review.school}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n4. Module Leader/Organiser\n'
        row.cells[1].text = f"\n{module.module_leader}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n5. Student numbers, achievement and progression\n'
        row.cells[1].text = f"\n{review.student_nap}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n6. Evaluation of the operation of the module\n'
        row.cells[1].text = f"\n{review.evolution_of_op_module}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n7. Evaluation of approach to teaching, assessment and feedback\n'
        row.cells[1].text = f"\n{review.evolution_to_teaching}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n8. Inclusive nature of the curriculum\n'
        row.cells[1].text = f"\n{review.inclusive_nature_of_curriculum}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n9. Effect of past changes\n'
        row.cells[1].text = f"\n{review.past_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n10. Proposed future changes\n'
        row.cells[1].text = f"\n{review.future_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\nOther comments\n'
        row.cells[1].text = f"\n{review.other_comment}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '12. Author and date'
        row.cells[1].text = f"\n{review.completed_by}\n{review.completion_date}\n"
        
        # Add a page break after each review
        document.add_page_break()
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={module.module_code}_Review.docx'
    document.save(response)

    return response

def export_ay_module_details(request, module_id, academic_year):
    module = get_object_or_404(Module, id=module_id)
    academic_year_obj = get_object_or_404(AcademicYear, academic_year=academic_year)
    reviews = ModuleReview.objects.filter(module=module, academic_year=academic_year_obj)
    
    document = Document()
    
    # Add logo image from URL
    logo_url = 'https://res.cloudinary.com/ddcynxxlr/image/upload/v1720968735/xypcmbpjm3szs6nwlh8e.jpg'
    response = requests.get(logo_url)
    if response.status_code == 200:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name

        header = document.sections[0].header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(tmp_file_path, width=Inches(2))  # Adjust size as needed
    
    for review in reviews:
        # Add heading for each review
        document.add_heading(f'Annual Module Quality Enhancement Report for {review.academic_year.academic_year}\n\n', level=1)
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '\nDescription\n'
        hdr_cells[1].text = '\nDetails\n'
        
        row = add_row_with_spacing(table)
        row.cells[0].text = '\n1. Module details\n'
        row.cells[1].text = f"\n{module.module_code} {module.module_name}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n2. Academic Year\n'
        row.cells[1].text = f"\n{review.academic_year.academic_year}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n3. School\n'
        row.cells[1].text = f"\n{review.school}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n4. Module Leader/Organiser\n'
        row.cells[1].text = f"\n{module.module_leader}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n5. Student numbers, achievement and progression\n'
        row.cells[1].text = f"\n{review.student_nap}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n6. Evaluation of the operation of the module\n'
        row.cells[1].text = f"\n{review.evolution_of_op_module}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n7. Evaluation of approach to teaching, assessment and feedback\n'
        row.cells[1].text = f"\n{review.evolution_to_teaching}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n8. Inclusive nature of the curriculum\n'
        row.cells[1].text = f"\n{review.inclusive_nature_of_curriculum}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n9. Effect of past changes\n'
        row.cells[1].text = f"\n{review.past_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n10. Proposed future changes\n'
        row.cells[1].text = f"\n{review.future_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\nOther comments\n'
        row.cells[1].text = f"\n{review.other_comment}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '12. Author and date'
        row.cells[1].text = f"\n{review.completed_by}\n{review.completion_date}\n"
        
        # Add a page break after each review
        document.add_page_break()
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={module.module_code}_{academic_year}_Review.docx'
    document.save(response)

    return response


# ======================Program Export ==================================================
def export_program_details(request, program_id):
    program = get_object_or_404(Program, id=program_id)
    reviews = ProgramReview.objects.filter(program=program).order_by('-academic_year')
    
    document = Document()

    # Add logo image from URL
    logo_url = 'https://res.cloudinary.com/ddcynxxlr/image/upload/v1720968735/xypcmbpjm3szs6nwlh8e.jpg'
    response = requests.get(logo_url)
    if response.status_code == 200:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name

        header = document.sections[0].header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(tmp_file_path, width=Inches(2))  # Adjust size as needed
    
    for review in reviews:
        # Add heading for each review
        document.add_heading(f'Annual Module Quality Enhancement Report for {review.academic_year.academic_year}\n\n', level=1)
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '\nDescription\n'
        hdr_cells[1].text = '\nDetails\n'
        
        row = add_row_with_spacing(table)
        row.cells[0].text = '\n1. Program details\n'
        row.cells[1].text = f"\n{program.program_code} {program.program_name}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n2. Academic Year\n'
        row.cells[1].text = f"\n{review.academic_year.academic_year}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n3. School\n'
        row.cells[1].text = f"\n{review.school}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n4. Program Leader/Organiser\n'
        row.cells[1].text = f"\n{program.program_leader}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n5. Student numbers, achievement and progression\n'
        row.cells[1].text = f"\n{review.student_nap}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n6. Evaluation of the operation of the program\n'
        row.cells[1].text = f"\n{review.evolution_of_op_program}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n7. Evaluation of approach to teaching, assessment and feedback\n'
        row.cells[1].text = f"\n{review.evolution_to_teaching}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n8. Inclusive nature of the curriculum\n'
        row.cells[1].text = f"\n{review.inclusive_nature_of_curriculum}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n9. Effect of past changes\n'
        row.cells[1].text = f"\n{review.past_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n10. Proposed future changes\n'
        row.cells[1].text = f"\n{review.future_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\nOther comments\n'
        row.cells[1].text = f"\n{review.other_comment}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '12. Author and date'
        row.cells[1].text = f"\n{review.completed_by}\n{review.completion_date}\n"
        
        # Add a page break after each review
        document.add_page_break()
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename={program.program_code}_Review.docx'
    document.save(response)

    return response


def export_ay_program_details(request, program_id, academic_year):
    program = get_object_or_404(Program, id=program_id)
    academic_year_obj = get_object_or_404(AcademicYear, academic_year=academic_year)
    reviews = ProgramReview.objects.filter(program=program, academic_year=academic_year_obj)
    
    document = Document()
    
    # Add logo image from URL
    logo_url = 'https://res.cloudinary.com/ddcynxxlr/image/upload/v1720968735/xypcmbpjm3szs6nwlh8e.jpg'
    response = requests.get(logo_url)
    if response.status_code == 200:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".jpg") as tmp_file:
            tmp_file.write(response.content)
            tmp_file_path = tmp_file.name

        header = document.sections[0].header
        header_paragraph = header.paragraphs[0]
        run = header_paragraph.add_run()
        run.add_picture(tmp_file_path, width=Inches(2))  # Adjust size as needed
    
    for review in reviews:
        # Add heading for each review
        document.add_heading(f'Annual Program Quality Enhancement Report for {review.academic_year.academic_year}\n\n', level=1)
        
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '\nDescription\n'
        hdr_cells[1].text = '\nDetails\n'
        
        row = add_row_with_spacing(table)
        row.cells[0].text = '\n1. Program details\n'
        row.cells[1].text = f"\n{program.program_code} {program.program_name}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n2. Academic Year\n'
        row.cells[1].text = f"\n{review.academic_year.academic_year}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n3. School\n'
        row.cells[1].text = f"\n{review.school}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n4. Program Leader/Organiser\n'
        row.cells[1].text = f"\n{program.program_leader}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n5. Student numbers, achievement and progression\n'
        row.cells[1].text = f"\n{review.student_nap}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n6. Evaluation of the operation of the program\n'
        row.cells[1].text = f"\n{review.evolution_of_op_program}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n7. Evaluation of approach to teaching, assessment and feedback\n'
        row.cells[1].text = f"\n{review.evolution_to_teaching}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n8. Inclusive nature of the curriculum\n'
        row.cells[1].text = f"\n{review.inclusive_nature_of_curriculum}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n9. Effect of past changes\n'
        row.cells[1].text = f"\n{review.past_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\n10. Proposed future changes\n'
        row.cells[1].text = f"\n{review.future_changes}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '\nOther comments\n'
        row.cells[1].text = f"\n{review.other_comment}\n"

        row = add_row_with_spacing(table)
        row.cells[0].text = '12. Author and date'
        row.cells[1].text = f"\n{review.completed_by}\n{review.completion_date}\n"
        
        # Add a page break after each review
        document.add_page_break()
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=Program_{program.program_code}_{academic_year}_Review.docx'
    document.save(response)

    return response

@api_view(['GET'])
def previous_year_module_reviews(request, module_code):
    module = get_object_or_404(Module, module_code=module_code)
    
    # Get the latest academic year
    latest_academic_year = AcademicYear.objects.order_by('-academic_year').first()
    
    # Get the previous academic year
    previous_academic_year = AcademicYear.objects.filter(academic_year__lt=latest_academic_year.academic_year).order_by('-academic_year').first()
    
    # Get the review for the previous academic year
    review = ModuleReview.objects.filter(module=module, academic_year=previous_academic_year).first()
    
    if review:
        review_data = {
            'academic_year': review.academic_year.academic_year,
            'school': review.school,
            'student_nap': review.student_nap,
            'evolution_of_op_module': review.evolution_of_op_module,
            'evolution_to_teaching': review.evolution_to_teaching,
            'inclusive_nature_of_curriculum': review.inclusive_nature_of_curriculum,
            'past_changes': review.past_changes,
            'future_changes': review.future_changes,
            'other_comment': review.other_comment,
            'completed_by': review.completed_by,
            'completion_date': review.completion_date.strftime('%Y-%m-%d') if review.completion_date else None,
        }
    else:
        review_data = None

    module_data = {
        'module_code': module.module_code,
        'module_name': module.module_name,
        'module_leader': module.module_leader,
    }

    return JsonResponse({
        'module': module_data,
        'review': review_data
    })